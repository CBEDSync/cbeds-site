# Builds the two Word versions of the admin docs. Word cannot render Mermaid, so the
# diagrams go in as PNGs from docs/diagrams. The .md files stay the source of truth -
# edit those, mirror the change here, and re-run:
#
#     pip install python-docx
#     python docs/make-docs.py
#
import os, re, struct
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
IMG = os.path.join(HERE, "diagrams")
OUT = os.path.dirname(HERE)                    # the repo root, next to the .md files

MAX_W, MAX_H = 6.9, 7.9          # inches of usable space at 0.8" margins on Letter


def png_size(path):
    with open(path, "rb") as f:
        head = f.read(24)
    return struct.unpack(">II", head[16:24])


def new_doc():
    d = Document()
    for s in d.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(0.8)
    n = d.styles["Normal"]
    n.font.name = "Segoe UI"
    n.font.size = Pt(10.5)
    n.paragraph_format.space_after = Pt(8)
    n.paragraph_format.line_spacing = 1.15
    return d


TOKEN = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)")


def rich(p, text):
    """Renders the small slice of Markdown these docs actually use."""
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**"):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x0F, 0x4C, 0x63)
        elif part.startswith("*"):
            p.add_run(part[1:-1]).italic = True
        else:
            p.add_run(part)
    return p


def para(d, text, **kw):
    return rich(d.add_paragraph(**kw), text)


def image(d, name):
    path = os.path.join(IMG, name + ".png")
    w, h = png_size(path)
    inches = min(MAX_W, MAX_H * w / h)
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    p.add_run().add_picture(path, width=Inches(inches))


def table(d, rows):
    t = d.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for a, b in rows:
        c = t.add_row().cells
        rich(c[0].paragraphs[0], a)
        rich(c[1].paragraphs[0], b)
        c[0].width = Inches(2.3)
        c[1].width = Inches(4.6)
    d.add_paragraph()
    return t


def footer(d, source):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("Generated from %s. Edit the Markdown, not this file - the "
                  "diagrams here are pictures and will go stale on their own." % source)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x6B, 0x7C, 0x84)


# ---------------------------------------------------------------- site doc
d = new_doc()
d.add_heading("How the site works", 0)
para(d, "Five HTML pages, no build step, no database, no server. Everything a visitor "
        "sees is a static file. The one exception is a single function that keeps an "
        "API key secret.")
image(d, "dataflow")
para(d, "**Edit the Excel -> run one `.bat` -> the site changes.** Nothing else touches "
        "the data.")

d.add_heading("The five pages", 1)
image(d, "pages")
para(d, "Only **CBEDSync** and **CBEDSense** load the data file - 143 KB gzipped. The "
        "other three are self-contained, so they stay light.")
para(d, "Each page opens with a full-screen hero built from the same `--hero-h` "
        "variable, with its own canvas or SVG artwork underneath the copy.")

d.add_heading("Where the data comes from", 1)
para(d, "`build.py` reads four sheets of `draft/CBEDSync.xlsx` and writes one file. It "
        "never writes back.")
image(d, "build")
para(d, "**`cbedsync-data.js` is generated - never edit it.** Every `.bat` overwrites it.")
para(d, "Fields are read **by position**, so a column inserted mid-sheet shifts "
        "everything after it and the build will publish a wrong graph without "
        "complaining. New columns go at the far right. (`Source` is the one exception - "
        "found by its heading.)")

d.add_heading("Publishing", 1)
image(d, "publish")
para(d, "Netlify redeploys on every push to `main`. There is no build command - it "
        "copies the files as they are.")

d.add_heading("The one piece of server code", 1)
para(d, '"Ask the graph" on CBEDSync answers from the graph in the browser. If an API '
        "key is set, `netlify/functions/ask.mjs` adds a short AI-written paragraph per "
        "section.")
image(d, "ask")
para(d, "The key lives only in Netlify's environment variables - **never in a page**, "
        "because every `.html` and `.js` file is public. Without a key the page falls "
        "back to its own answers, so nothing breaks.")

d.add_heading("Submissions", 1)
para(d, "Two forms collect entries from outside the team. Nothing they send reaches the "
        "site without a person copying it into the Excel - see **Handling "
        "submissions**.")

d.add_heading("Files that matter", 1)
table(d, [
    ("`draft/CBEDSync.xlsx`", "**the master data - edit this**"),
    ("`cbedsync-data.js`", "generated, 143 KB gzipped, don't edit"),
    ("`build.py`", "Excel to data file"),
    ("`*.bat`", "rebuild / publish / collect submissions"),
    ("`cbeds-assistant.js`", 'shared "ask the graph" engine (CBEDSync + CBEDSense)'),
    ("`netlify/functions/ask.mjs`", "the only server code; holds no key"),
    ("`.env`", "your local keys - git-ignored, never published"),
])
para(d, "Setup and hosting: `HOW-TO-HOST-AND-UPDATE.md`.")
footer(d, "HOW-THE-SITE-WORKS.md")
d.save(os.path.join(OUT, "How the site works.docx"))

# --------------------------------------------------------- submissions doc
d = new_doc()
d.add_heading("Handling submissions", 0)
para(d, "Two forms let people outside the team put entries forward: **Share your work** "
        "(CBEDSense) and the **Alliance Charter** (CBEDSynergy).")
para(d, "**Nothing they send reaches the site on its own.** It only gets published if "
        "you copy it into `draft/CBEDSync.xlsx` yourself.")
image(d, "submissions")

d.add_heading("Setup (once)", 1)
para(d, "At https://app.netlify.com/projects/cbeds/forms")
for i, t in enumerate([
    "Turn **form detection** on. Two forms appear: `cbeds-charter`, "
    "`cbedsense-submission`.",
    "Add an **email notification** for each - *before* testing, or you get no email.",
    "Make a token at https://app.netlify.com/user/applications, then paste it into "
    "`.env` as `NETLIFY_TOKEN=...` (git-ignored; if Notepad saves `.env.txt`, set "
    "*Save as type: All Files*).",
], 1):
    para(d, t, style="List Number")

d.add_heading("Reading the staging file", 1)
table(d, [
    ("Columns A onwards", "The master's own columns. `Source` already says `Public`."),
    ("**Amber cell, row 1**", "Says exactly what to copy, e.g. `<- copy A:BM only`. "
                              "**Read it, don't remember it** - it moves if the master "
                              "gains or loses a column."),
    ("Green `Review:` columns", "Your notes. Never copied across."),
    ("**Red text**", "A researched suggestion, *not* what the submitter wrote. Check "
                     "before approving."),
])

d.add_heading("The research step", 1)
para(d, "Ask Claude, pasting what the run printed. It reads a copy of `CBEDSync.xlsx`, "
        "searches online, and fills in type, dates, location, themes and links - all "
        "in red.")
para(d, "Two rules it follows: links only ever point at entities **already in the "
        "graph** (an invented name makes a dead link, not an error), and every "
        "suggestion is explained in `Review: Notes`. It never overwrites anything the "
        "submitter typed.")

d.add_heading("Watch out", 1)
for t in [
    "**Never insert a column mid-sheet** in Agent / Project / Output. The build reads "
    "by position - it will publish a wrong graph without complaining. New columns go "
    "at the far right.",
    "**`Source` can't be reconstructed later.** It comes across as `Public` already; "
    "don't clear it.",
    "**A Charter signature is two things.** The organisation becomes an Agent row; the "
    "named lead, date and commitments have no home there - keep those with your "
    "signatory record.",
    "`draft/Submissions.xlsx` holds names and emails and is **kept out of GitHub**.",
    "Free Netlify allows ~100 submissions/month (*Billing > Usage*).",
]:
    para(d, t, style="List Bullet")

d.add_heading("If it complains", 1)
table(d, [
    ("`No NETLIFY_TOKEN found`", "Step 3 skipped, or Notepad saved `.env.txt`"),
    ("`Netlify refused the token`", "Token wrong or revoked - make a new one"),
    ("`Could not find a Netlify project...`", "It lists what your token can see; set "
                                              "`NETLIFY_SITE`"),
    ("`No CBEDS forms on cbeds`", "Detection off, or no deploy has carried the forms"),
    ("`...does not match the master's columns`", "Delete `draft/Submissions.xlsx`, run "
                                                 "again (copy out any `Review:` notes "
                                                 "first)"),
    ('`no "Source" column on ...`', "Add `Source` at the far right of that sheet"),
    ("`names shared by more than one entry`", "Not an error - two entries share a name, "
                                              "usually a company and its product"),
    ("`Python was not found`", "See `HOW-TO-HOST-AND-UPDATE.md`"),
])

d.add_heading("Files", 1)
table(d, [
    ("`get-submissions.bat`", "collect"),
    ("`rebuild-only.bat`", "rebuild and stop"),
    ("`update-website.bat`", "publish"),
    ("`draft/CBEDSync.xlsx`", "**the master**"),
    ("`draft/Submissions.xlsx`", "the waiting room (generated, safe to delete)"),
    ("`.env`", "your token"),
])
para(d, "Hosting and the Excel itself: `HOW-TO-HOST-AND-UPDATE.md`.")
footer(d, "SUBMISSIONS.md")
d.save(os.path.join(OUT, "Handling submissions.docx"))

print("wrote both files to", OUT)
